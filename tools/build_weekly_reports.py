#!/usr/bin/env python3
"""Build four evidence-grounded personal weekly reports as DOCX files."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "week1-2-weekly-reports"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
BLACK = RGBColor(0, 0, 0)
BODY_FONT = "Arial"
EAST_ASIA_FONT = "Arial Unicode MS"


REPORTS = [
    {
        "filename": "张家豪-第1周个人周报.docx",
        "name": "张家豪",
        "week": "第 1 周",
        "period": "2026 年 7 月 6 日 - 7 月 10 日",
        "confirmation": "已按本次重做过程核对",
        "sections": [
            ("一、组号", "20"),
            ("二、实训题目", "面向 CRAG-MM 的多模态多轮检索增强问答系统设计与实现"),
            ("三、本周小组分工", "小组按相对均衡的技术职责推进：张家豪负责 Agent 总体架构、模型生成后端、Prompt、答案约束与任务1方案；俞凡负责检索接口、KG/Web 数据解析、证据排序、评测脚本与实验记录。环境复现、固定样例确认、任务2联调和材料复核由两人共同完成。"),
            ("四、本人负责的部分", "本人负责把课题从“完成一个能回答的程序”细化为可实际验收的技术路线。具体包括梳理 CRAG-MM 三项任务的输入、检索源、输出格式、评分方式和运行限制，并确定“纯视觉 baseline → 单源图像 KG → 图像 KG 与网页多源融合”的递进路线。我还检查了数据样本与搜索接口所能提供的信息，为 Agent 定下“先结构化证据，再生成答案”的设计原则。在提示语方面，我负责划分当前图片的可见事实、图像知识库信息和网页信息，防止模型把相似图片的内容误当为当前图片事实。同时确定本机优先调试、无法满足真实评测时转用云端 GPU 的边界。"),
            ("五、工作完成情况、问题与反思", "本周完成了课程安排、starter kit、数据集、检索 API、Agent 接口和提交规范的阅读，明确了每轮回答需尽量控制在 10 秒内、答案不超过 75 BPE token、需保持离线运行与提交格式兼容等约束。项目已建立 Python 3.11 独立环境，并确认 torch、datasets、transformers、cragmm-search-pipeline 和 MLX-VLM 均可导入，为后续实际运行打下基础。本周最重要的反思是：早期交付中有代码和说明，但没有能追溯到具体样例的运行证据，因此不能仅根据文档就判定完成。后续我将完成依据统一为固定样例、配置文件、原始 JSONL trace、测试结果和渲染预览。代码实现使用了 AI 辅助，本人负责目标设定、技术取舍与结果核对。"),
        ],
    },
    {
        "filename": "张家豪-第2周个人周报.docx",
        "name": "张家豪",
        "week": "第 2 周",
        "period": "提前完成稿：2026 年 7 月 12 日；计划周为 7 月 13 日 - 7 月 17 日",
        "confirmation": "已按本次重做过程核对",
        "sections": [
            ("一、组号", "20"),
            ("二、实训题目", "面向 CRAG-MM 的多模态多轮检索增强问答系统设计与实现"),
            ("三、本周小组分工", "张家豪负责 Agent V2 架构、Vision/Task1/Task2 模式隔离、生成后端适配、Prompt、拒答和答案后处理；俞凡负责检索证据结构、字段排序、实验脚本、指标与错误案例。两人共同确认固定样例、真实运行条件和中期汇报内容。"),
            ("四、本人负责的部分", "本人重点负责将旧原型重构为后端无关的统一 Agent。在架构上，本机通过 MLX-VLM 完成调试，官方 CUDA 环境通过 vLLM 运行，但两者共用相同的检索、证据组装和答案后处理逻辑。我设计了显式 task_mode，使 Vision 不检索、Task1 在代码层面禁止网页检索、Task2 才能加入网页证据。在生成策略上，我将当前图片可见内容、相似图片 KG 信息和网页证据分开呈示，要求模型只使用能直接支持答案的信息。同时加入证据不足、证据冲突和检索失败三类拒答边界，以及输入长度、后端输出数量和 75-token 截断校验。"),
            ("五、工作完成情况、问题与反思", "本周已完成 V2 Agent、正式提交入口、固定抽样与汇总脚本、环境快照和 19 项自动化测试。用可控的检索结果运行了 120 次流程验证：Task1 没有引入网页证据，Task2 则能把网页证据编号并传入回答环节。同时，项目已下载真实 validation 分片（387 条）并固定抽取 30 条样例，也完成了真实图片的本地模型冒烟。小模型在 3.113 秒内将 Kadhi Chawal 误认为 Chicken curry，因此只能证明推理链路运行，不能证明效果达标。完整网页检索索引约 9.5GB，还需图像索引与 11B 视觉模型，因此真实 Task1/Task2 的 30+30 对照仍需云端 GPU 完成。我会将工程流程正确与真实性能验证明确分开记录。"),
        ],
    },
    {
        "filename": "俞凡-第1周个人周报-待本人确认.docx",
        "name": "俞凡",
        "week": "第 1 周",
        "period": "2026 年 7 月 6 日 - 7 月 10 日",
        "confirmation": "待俞凡本人确认后提交",
        "sections": [
            ("一、组号", "20"),
            ("二、实训题目", "面向 CRAG-MM 的多模态多轮检索增强问答系统设计与实现"),
            ("三、本周小组分工", "小组按相对均衡的技术职责推进：张家豪负责 Agent 总体架构、模型生成后端、Prompt、答案约束与任务1方案；俞凡负责检索接口、KG/Web 数据解析、证据排序、评测脚本与实验记录。环境复现、固定样例确认、任务2联调和材料复核由两人共同完成。"),
            ("四、本人负责的部分", "本周分配给本人的工作是理解并梳理项目的检索与评测基础。具体需检查项目依赖、数据集结构、图像检索接口和本地评测流程，并记录 search pipeline 的初始化方式、image search 返回的实体名称、检索分数和属性字段。针对数据集，需明确 validation split 中 interaction_id、问题、标准答案、问题类别和领域等字段的含义，为后续固定抽样和人工复核打下基础。我还需为实验记录检查必备字段，包括运行模式、查询内容、检索结果、最终证据、回答、耗时与异常类型。以上为本人分工与待本人确认的工作范围，不以 AI 生成内容代替个人实际核查。"),
            ("五、工作完成情况、问题与反思", "从项目现有证据看，已建立可复现的 Python 3.11 环境，关键依赖导入检查正常，并已从真实 single-turn validation parquet 分片读取 387 条数据。固定样例采用种子 20260712，从中选取 30 条，已保存 interaction_id、问题类别、领域、问题和标准答案。这些记录使后续不同版本能在同一批数据上比较，而不是每次随意抽题。当前问题是完整图像与网页检索索引体积较大，且网络下载不稳定，所以还没有得出真实的 Task1/Task2 成绩。本人提交前需复核运行命令、样例清单和自己实际参与的分析内容；未经确认的事项不得写为个人已完成工作。"),
        ],
    },
    {
        "filename": "俞凡-第2周个人周报-待本人确认.docx",
        "name": "俞凡",
        "week": "第 2 周",
        "period": "提前完成稿：2026 年 7 月 12 日；计划周为 7 月 13 日 - 7 月 17 日",
        "confirmation": "待俞凡本人确认后提交",
        "sections": [
            ("一、组号", "20"),
            ("二、实训题目", "面向 CRAG-MM 的多模态多轮检索增强问答系统设计与实现"),
            ("三、本周小组分工", "张家豪负责 Agent V2 架构、生成后端、Prompt、拒答和答案后处理；俞凡负责 image KG 与网页结果解析、问题相关字段排序、实验运行、指标统计和错误案例。两人共同进行任务2联调、中期材料复核和云端实验准备。"),
            ("四、本人负责的部分", "本周分配给本人的工作是完善检索与评测链路的证据处理部分。对图像检索结果，需将 KG 实体名称、检索分数和属性字段解析为结构化证据，再根据问题关键词对属性进行排序，避免把不相关的字段全部提供给模型。对网页检索结果，需检查标题、摘要、URL 和分数，完成空摘要过滤、URL 去重、长度控制和证据编号。评测方面，需准备固定抽样、原始 JSONL/CSV、耗时统计、拒答率和错误分析表，并在任务1 中确认网页检索不会被调用。以上为分工范围，提交前仍需由本人核对实际负责和运行的部分。"),
            ("五、工作完成情况、问题与反思", "从项目现有证据看，结构化证据解析、问题相关字段排序、网页去重与阈值过滤、检索异常 trace，以及可重复的实验和汇总脚本都已进入项目。用可控的检索器和模型生成了 120 条 trace：任务1 没有网页证据，任务2 的网页证据路径按设计工作。这只说明程序流程符合设计，不等于 CRAG-MM 真实性能。后续还需在完整图像与网页索引上运行 30+30 固定样例，分别整理成功、失败、合理拒答、网页噪声和跨源冲突案例。本人提交前应确认自己实际执行过的命令、复核过的数据与撰写过的分析，不应照抄未参与的工作。"),
        ],
    },
]


def set_run_font(run, size=None, color=BLACK, bold=None, italic=None):
    run.font.name = BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), BODY_FONT)
    rpr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rpr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.05
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        # Named override: compact_weekly_two_page keeps expanded reflections readable.
        ("Heading 2", 13, BLUE, 8, 4),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def build_report(report):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header_run = header.add_run("智能技术实训｜个人周报")
    set_run_font(header_run, size=9, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run(f"{report['week']}个人周报"), size=23, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(subtitle.add_run("面向 CRAG-MM 的多模态多轮检索增强问答系统"), size=13, color=MUTED)

    for label, value in [
        ("姓名", report["name"]),
        ("组号", "20"),
        ("周期", report["period"]),
        ("事实确认", report["confirmation"]),
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(f"{label}："), size=10.5, bold=True)
        set_run_font(paragraph.add_run(value), size=10.5, color=MUTED if label == "事实确认" else BLACK)

    for heading, body in report["sections"]:
        # The expanded reflection is deliberately kept as a readable second-page unit.
        if heading.startswith("五、"):
            doc.add_page_break()
        doc.add_paragraph(heading, style="Heading 2")
        paragraph = doc.add_paragraph()
        set_run_font(paragraph.add_run(body), size=10.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(
        note.add_run("说明：本周报按项目现有运行证据撰写；合成契约测试不等同于 CRAG-MM 官方成绩。"),
        size=9,
        color=MUTED,
        italic=True,
    )

    doc.core_properties.title = f"{report['name']} {report['week']}个人周报"
    doc.core_properties.subject = "智能技术实训｜CRAG-MM"
    doc.core_properties.author = report["name"]

    OUTPUT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT / report["filename"]
    doc.save(path)
    return path


def main():
    for report in REPORTS:
        print(build_report(report))


if __name__ == "__main__":
    main()
